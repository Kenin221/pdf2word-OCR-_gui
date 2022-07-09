from pathlib import Path
import tkinter as tk
from PIL import ImageTk, Image
from tkinter import filedialog, ttk
import os
from docx import Document
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import time

ins = 0
win = tk.Tk()
win.geometry('1280x720')
win.title('PDF轉WORD')
win.iconbitmap('wm','icon.ico')
bg = tk.Canvas(win, width=1280,height=720)
imagepath = 'bg.png'
img = Image.open(imagepath)
photo = ImageTk.PhotoImage(img)
bg.create_image(640,360,image=photo)
bg.pack()

def p2_path_onefile():
    global file_path
    file_path = filedialog.askopenfilename(parent=win, initialdir='~/Desktop', filetypes=(('PDF files','*pdf'),('PDF files','*pdf'))) #選擇檔案(initialdir為指定起始目錄)
    main()
def close():
    for widgets in win.winfo_children():
            widgets.destroy()
def back_mode():
    close() 
    global ins, bg ,imagepath , img , photo
    ins = 0
    bg = tk.Canvas(win, width=1280,height=720)  
    imagepath = 'bg.png'
    img = Image.open(imagepath)
    photo = ImageTk.PhotoImage(img)
    bg.create_image(640,360,image=photo)
    bg.pack()
    lb1 = tk.Label(text='PDF轉WORD',fg='blue',font=('stick',30),bg='#FFFAF0')
    lb1.place(x=50,y=25)
    but1 = tk.Button(text='選擇檔案',fg='green',font=('setofont',20),command=p2_path_onefile,bg='#FFFAF0')
    but1.place(x=100,y=130)

ins = 0
lb1 = tk.Label(text='PDF轉WORD',fg='blue',font=('stick',30),bg='#FFFAF0')
lb1.place(x=50,y=25)
but1 = tk.Button(text='選擇檔案',fg='green',font=('setofont',20),command=p2_path_onefile,bg='#FFFAF0')
but1.place(x=100,y=130)

def main():
    close()
    try:    
        bg = tk.Canvas(win, width=1279,height=720)
        imagepath = 'bg2.png'
        img = Image.open(imagepath)
        photo = ImageTk.PhotoImage(img)
        bg.create_image(640,360,image=photo)
        bg.pack()
        load_label = tk.Label(win, text="目前進度 : 0.00%", fg='yellow', font=('stick','30'),bg='#F5B6A7')
        load_block = ttk.Progressbar(win, length=1000, mode='determinate', orient=tk.HORIZONTAL)
        load_label.place(x=500,y=100)
        load_block.place(x=150,y=50)
        name = str(Path(file_path).stem)
        path = os.path.split(file_path)
        file = path[0]
        document = Document()
        images = convert_from_path(file_path,poppler_path=r'poppler-22.04.0/Library/bin')
        win.update()
        load_block['value'] += 15
        load_label['text'] = "目前進度 : 15.0%"
        win.update()
        time.sleep(0.4)
        try:
            if count > 1:
                for i in range(count):
                    os.remove('temp'+str(i)+'.png')
            elif count == 1:
                os.remove('temp0.png')
        except:pass
        load_block['value'] += 35
        load_label['text'] = "目前進度 : 50.0%"
        win.update()
        time.sleep(0.3)
        count = len(images)
        for i in range(count):
            images[i].save('temp'+str(i)+'.png', 'PNG')

        pytesseract.pytesseract.tesseract_cmd = r"Tesseract-OCR/tesseract.exe"
        load_block['value'] += 15
        load_label['text'] = "目前進度 : 65.0%"
        win.update()
        time.sleep(0.2)
        if count > 1:
            for i in range(count):
                img_name = 'temp'+str(i)+'.png'
                img = Image.open(img_name)
                text = pytesseract.image_to_string(img, lang='chi_tra')
                globals()['p'+str(i)] = document.add_paragraph('')
                globals()['p'+str(i)].add_run(text)
                document.save(file + '/' +  name + '.docx')
        elif count == 1:
            img_name = 'temp0.png'
            img = Image.open(img_name)
            text = pytesseract.image_to_string(img, lang='chi_tra')
            p1 = document.add_paragraph('')
            p1.add_run(text)
            document.save(file + '/' +  name + '.docx')
        load_block['value'] += 25
        load_label['text'] = "目前進度 : 90.0%"
        win.update()
        time.sleep(0.2)
        try:
            if count > 1:
                for i in range(count):
                    os.remove('temp'+str(i)+'.png')
            elif count == 1:
                os.remove('temp0.png')
        except:pass
    except:
        if  file_path ==  '':
            close()
            bg = tk.Canvas(win, width=1280,height=1280)
            imagepath = 'bad.png'
            img = Image.open(imagepath)
            photo = ImageTk.PhotoImage(img)
            bg.create_image(640,360,image=photo)
            bg.pack()
            load_label = tk.Label(win, text="目前進度 : 0.00%", fg='yellow', font=('stick','30'),bg='#F5B6A7')
            load_block = ttk.Progressbar(win, length=1000, mode='determinate', orient=tk.HORIZONTAL)
            load_label.place(x=430,y=80)
            load_block.place(x=150,y=25)
            load_label['text'] = '未選擇文件，將返回首頁!'
            win.update()
            time.sleep(2)
            global ins
            ins = 1
            time.sleep(2)
            return back_mode()
    if ins == 0:
        load_block['value'] += 10
        load_label['text'] = "目前進度 : 100.0%"
        win.update()
        time.sleep(0.2)
        close()
        bg = tk.Canvas(win, width=1280,height=1280)
        imagepath = 'yes.png'
        img = Image.open(imagepath)
        photo = ImageTk.PhotoImage(img)
        bg.create_image(640,360,image=photo)
        bg.pack()
        load_label = tk.Label(win, text="目前進度 : 100.00%", fg='yellow', font=('stick','30'),bg='#F5B6A7')
        load_block = ttk.Progressbar(win, length=1000, mode='determinate', orient=tk.HORIZONTAL)
        load_block['value'] += 100
        load_label.place(x=500,y=100)
        load_block.place(x=150,y=50)
        load_label.place(x=200,y=100)
        load_label['text'] = '轉換成功!word檔將會放在與PDF相同的資料夾中'
        win.update()
        time.sleep(0.35)
        back = tk.Button(text='回首頁',font=('setofont',40),command=back_mode)
        back.place(x=1050,y=600)
        return win.update()
win.mainloop()